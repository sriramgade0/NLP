from paddleocr import PaddleOCR
import re
from docx import Document
from PIL import Image
import matplotlib.pyplot as plt

# Initialize PaddleOCR
ocr = PaddleOCR(use_angle_cls=True, lang='en')

# Load Image
image_path = 'potato_chips_2.png'  # Update with your image path
image = Image.open(image_path)

# Show the image (optional)
plt.imshow(image)
plt.axis('off')
plt.show()

# Perform OCR
result = ocr.ocr(image_path)

# Extract text from OCR result
ocr_text = "\n".join([line[1][0] for line in result[0]])
print("Extracted OCR Text:")
print(ocr_text)

# Save OCR text into a Word document
document = Document()
document.add_heading('OCR Extracted Text', level=1)
document.add_paragraph(ocr_text)

# Save extracted metadata
patterns = {
    "Name": r"(?i)(name|product):?\s*(.+)",
    "Weight or Volume": r"(?i)(weight|volume):?\s*(\d+[\w\s]+)",
    "Nutrition Information": r"(?i)(nutrition(?:al)? information):?\s*((?:.|\n)*?)(?:ingredients|$)",
    "Ingredients": r"(?i)(ingredients):?\s*((?:.|\n)*?)$"
}

metadata = {}
for key, pattern in patterns.items():
    match = re.search(pattern, ocr_text)
    if match:
        metadata[key] = match.group(2).strip()

# Add extracted metadata to the document
document.add_heading('Extracted Metadata', level=1)
for key, value in metadata.items():
    document.add_paragraph(f"{key}: {value}")

# Save the document
output_path = 'OCR_Extracted_Text.docx'
document.save(output_path)
print(f"OCR text and metadata saved to {output_path}")
